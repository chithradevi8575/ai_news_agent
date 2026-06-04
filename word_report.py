from docx import Document

def create_doc(news, ai_text):
    doc = Document()
    doc.add_heading("🤖 AI NEWS REPORT", 0)

    doc.add_heading("📰 TOP 10 HEADLINES", level=1)

    for i, item in enumerate(news, 1):
        doc.add_paragraph(f"{i}. {item['title']}")
        doc.add_paragraph(f"Link: {item['link']}")
        doc.add_paragraph("")

    doc.add_heading("📌 OVERALL SUMMARY & HIGHLIGHTS", level=1)
    doc.add_paragraph(ai_text)

    file_name = "AI_News_Report.docx"
    doc.save(file_name)

    return file_name